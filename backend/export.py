from docx import Document
from datetime import datetime

def export_word(params, output, image):

    doc = Document('assets/report_template.docx')

    if output[4]:
        volume = output[3]
        unit_volume = 'litre'
        mass = params['volume']
        unit_mass = params['unit_volume']
    else:
        volume = params['volume']
        unit_volume = params['unit_volume']
        mass = round(output[5],2)
        unit_mass = 'kg'

    placers = {
        '{system}': params['system_name'],
        '{medium}': params['medium'],
        '{volume}': str(volume),
        '{unit_volume}': unit_volume,
        '{mass}': str(mass),
        '{unit_mass}': unit_mass,
        '{leak_rate_refprop}': str(output[1]),
        '{leak_rate_ideal}': str(output[0]),
        '{bubble_rate}': str(output[2]),
        '{current_date}': datetime.now().strftime('%d-%m-%Y'),
        '{start_time}': datetime.strptime(params['start_time'], '%Y-%m-%dT%H:%M').strftime('%d-%m-%Y %H:%M:%S'),
        '{end_time}': datetime.strptime(params['end_time'], '%Y-%m-%dT%H:%M').strftime('%d-%m-%Y %H:%M:%S')
    }


    def replace_text(paragraphs):
        for paragraph in paragraphs:
            for placeholder, value in placers.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

    replace_text(doc.paragraphs)

    for section in doc.sections:
        page_width = section.page_width - section.left_margin - section.right_margin
        for table in section.first_page_header.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text(cell.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placers.items():
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(placeholder, value)
                            paragraph.text = paragraph.text.replace(placeholder, value)

    doc.add_picture(image, width=page_width )

    return doc
    #doc.save('assets/changed_report.docx')