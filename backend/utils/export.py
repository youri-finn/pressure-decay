from docx import Document
from datetime import datetime, timedelta
import os

APP_ENV = os.getenv('APP_ENV', 'local')


def export_word(parameters, results, images):

    if APP_ENV == 'local':
        template_path = 'assets/report_template.docx'
    else:
        template_path = '/opt/app-root/src/backend/assets/report_template.docx'

    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        raise FileNotFoundError(f"Word template not found at specified path")

    if results['measured_volume']:
        volume = results['measured_volume']
        unit_volume = 'liter'
        mass = parameters['mass']
        unit_mass = parameters['unit_volume']
    else:
        volume = parameters['volume']
        unit_volume = parameters['unit_volume']
        mass = round(results['trendline_parameters_mass'][1], 2)
        unit_mass = 'kg'

    periodic_limit = 24  # hours (be careful, as this is also temporarily hardcoded in the analysis file)

    if parameters['periodic_limit_off']:
        end_time = datetime.strptime(parameters['end_time'], '%Y-%m-%dT%H:%M:%S').strftime('%d-%m-%Y %H:%M:%S')
    else:
        end_time = (datetime.strptime(parameters['start_time'], '%Y-%m-%dT%H:%M:%S')
                    + timedelta(hours=periodic_limit)*int(results['periods'])).strftime('%d-%m-%Y %H:%M:%S')

    placers = {
        '{system}': parameters['system_name'],
        '{medium}': parameters['medium'],
        '{volume}': str(volume),
        '{unit_volume}': unit_volume,
        '{mass}': str(mass),
        '{unit_mass}': unit_mass,
        '{leak_rate_refprop}': str(results['mass_rate']),
        '{leak_rate_ideal}': str(results['ideal_gas_rate']),
        '{bubble_rate}': str(results['bubble_rate']),
        '{current_date}': datetime.now().strftime('%d-%m-%Y'),
        '{stabilization_time}': str(results['stabilization_time']),
        '{total_test_time}': str(results['total_test_time']),
        '{start_time}': datetime.strptime(parameters['start_time'], '%Y-%m-%dT%H:%M:%S').strftime('%d-%m-%Y %H:%M:%S'),
        '{end_time}': end_time
    }

    target_starts = ['Leak rate (mass-density)', 'Leak rate (ideal gas law)', 'Bubble diameter rate']

    def replace_text(paragraphs):
        for paragraph in paragraphs:
            for placeholder, value in placers.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

    replace_text(doc.paragraphs)

    for paragraph in doc.paragraphs:
        for start in target_starts:
            if paragraph.text.startswith(start):
                prefix, suffix = paragraph.text.split(':')
                paragraph.clear()
                paragraph.add_run(prefix + ':')
                paragraph.add_run(suffix).bold = True

    for section in doc.sections:
        page_width = section.page_width - section.left_margin - section.right_margin
        for table in section.first_page_header.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text(cell.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placers.items():
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(placeholder, value)
                            paragraph.text = paragraph.text.replace(placeholder, value)

    for image in images:
        doc.add_picture(image, width=page_width )

    return doc
